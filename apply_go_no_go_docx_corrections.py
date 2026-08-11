#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Apply targeted go/no-go manuscript corrections to final DOCX files."""

from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


ROOT_CLEAN = Path("sent-foods-4487055_major_revision_r2_r3_FINAL_CLEAN.docx")
ROOT_HIGHLIGHTED = Path("sent-foods-4487055_major_revision_r2_r3_FINAL_HIGHLIGHTED.docx")
PKG_HIGHLIGHTED = Path("SUBMISSION_PACKAGE/Manuscripts/01_Manuscript_FINAL_HIGHLIGHTED.docx")
PKG_CLEAN = Path("SUBMISSION_PACKAGE/Manuscripts/02_Manuscript_FINAL_CLEAN.docx")

TITLE = "In-Bag Pressure and Temperature Monitoring of Palletised 3-L Bag-in-Box Wine under Static Chamber Conditions"


PARAGRAPH_UPDATES = {
    1: TITLE,
    11: (
        "Bag-in-Box (BiB) packaging can reduce the logistical and environmental burden of wine exports, but its "
        "behaviour under combined stacking and thermal stress is still poorly documented. This study monitored "
        "3-L BiB units during a 20-day static chamber trial in two chambers nominally set to 19 °C and 50 °C, "
        "followed by a 15-day single-stack verification trial under a nominal 50 °C condition. In the main trial, "
        "the nominal 50 °C chamber did not produce 50 °C wine exposure. During the late storage window, in-bag "
        "temperatures averaged 34.18 ± 0.36 °C in top units and 26.08 ± 0.23 °C in bottom units. In the nominal "
        "19 °C chamber, top and bottom units averaged 23.50 ± 0.21 °C and 22.84 ± 0.30 °C, respectively. Internal "
        "pressure traces shared a broad common-mode pattern, and no barometric reference was logged. In the "
        "uncorrected exploratory sensor-level analysis, bottom-position BiBs had a larger peak baseline-referred "
        "pressure increase than top-position BiBs (adjusted difference 20.4 mbar; 95% CI 4.8-35.9 mbar; "
        "F(1,8) = 9.14; p = 0.017), but the position term for each sensor's maximum common-mode residual was not "
        "statistically detectable (F(1,8) = 3.21; p = 0.111). The pressure findings are therefore interpreted as "
        "transient position-related offsets superimposed on common-mode absolute-pressure variation. Descriptive "
        "bulk chemistry summaries for wine from the nominally warmer chamber showed lower free and total SO₂ and "
        "higher volatile acidity after 20 days, but chemical records were not resolved by stack position and do "
        "not support treatment-level inference. Lactic acid was numerically higher, but raw chemical replicates "
        "and microbiological data were unavailable. In the verification trial, sensor maxima ranged from 45.0 to "
        "48.8 °C and no in-bag sensor reached 50 °C. The results support selected sentinel sensing to document "
        "actual package conditions, while highlighting the limits of inferring wine exposure from chamber set "
        "points alone."
    ),
    19: (
        "Previous studies have examined storage temperature and packaging effects on wine composition, including "
        "sulphur dioxide loss, volatile acidity and sensory ageing [17,20,21]. Most of that evidence concerns "
        "bottled wines, static storage, or chemical composition alone. Less is known about the pressure developed "
        "inside palletised BiB units during static chamber storage intended to approximate selected logistics "
        "stresses, or about how measured in-package thermal exposure compares with nominal chamber settings."
    ),
    20: (
        "The present work was designed to test whether stack position and a severe nominal chamber condition "
        "measurably affected internal pressure and selected chemical parameters of commercial 3-L BiB white wine "
        "under static chamber storage. The original intention was to compare a neutral reference condition with "
        "a nominal 50 °C high-temperature condition. Because the main trial did not achieve 50 °C in-bag exposure, "
        "the realised study is interpreted as a comparison between nominal chamber conditions and measured package "
        "temperatures, supplemented by a single-stack verification trial that approached the set point more closely "
        "but still did not reach 50 °C inside every bag."
    ),
    22: (
        "The purpose of the work was to investigate selected storage stresses that wine packages may encounter "
        "during export logistics under static chamber conditions. The study did not reproduce vibration, humidity, "
        "diurnal thermal cycling or a full container thermal profile. It comprised a 20-day main trial and a "
        "15-day verification trial. In the main trial, palletised units were stored in parallel in two "
        "temperature-controlled chambers defined by their nominal set points:"
    ),
    24: (
        "Condition 2: chamber nominally set at 50 °C, selected as a deliberately severe nominal set point relevant "
        "to worst-case non-refrigerated storage scenarios rather than as a routine shipment temperature [15]."
    ),
    28: (
        "The available experimental records did not include the chamber model, chamber volume, fan power, airflow "
        "rate, or a formal pallet-occupancy ratio. These unrecorded chamber characteristics limit diagnosis of "
        "the thermal mismatch and vertical temperature gradient observed in the main trial. The manuscript "
        "therefore reports only measured in-bag temperatures and treats chamber-performance explanations as "
        "hypotheses rather than established mechanisms."
    ),
    39: (
        "Free sulphur dioxide, total sulphur dioxide, titratable acidity (g/L tartaric acid), volatile acidity "
        "(g/L acetic acid), ethanol, malic acid, lactic acid, total phenols and pH were determined according to "
        "the protocols described in a previous study [18]. The available study records report chemical values as "
        "triplicate determinations for main-trial wine grouped by nominal chamber condition, but they do not show "
        "whether the triplicates were independent package-level samples or analytical repeat measurements. "
        "Chemical sampling was not resolved by stack position, and no microbiological measurements or viable cell "
        "counts were performed. The available pressure/temperature records did not include raw chemical replicate "
        "values."
    ),
    48: (
        "Chemical results were not used for treatment-level inference in the revised manuscript. The available "
        "study records report triplicate determinations by nominal chamber condition, but raw replicate records "
        "were not retained and the independence of those triplicates could not be verified. Therefore, chemical "
        "parameters are presented descriptively as mean ± SD, without ANOVA, Tukey letters or p-values. The "
        "lactic-acid SD for the nominal 50 °C chamber was traced to the originally submitted Table 2 "
        "(0.29 ± 0.02 g/L); no primary experimental record supported the alternative 0.20 g/L SD."
    ),
    57: (
        "The chemical comparison concerns wine from the main trial grouped by nominal chamber condition. The labels "
        "19 °C and 50 °C in Table 2 denote chamber set points, not measured wine temperatures. Late in the main "
        "trial, in-bag sensors showed that the nominal 50 °C chamber produced 34.18 ± 0.36 °C at top sensors and "
        "26.08 ± 0.23 °C at bottom sensors. Packages in the nominal 19 °C chamber were about 23 °C (Section 3.2). "
        "The chemical results should therefore be read as descriptive differences associated with the nominally "
        "warmer chamber condition, not as the effect of exposing wine to 50 °C."
    ),
    58: (
        "Lactic acid was numerically higher in the wine from the nominally warmer chamber (0.29 ± 0.02 g/L) than "
        "in the reference (0.16 ± 0.05 g/L) (Table 2). This difference is not treated as statistically verified "
        "because raw chemical replicates were unavailable. The 0.29 ± 0.02 g/L value is reported because it is "
        "traceable to the originally submitted Table 2; no primary experimental record supports the alternative "
        "0.20 g/L SD. The available measurements do not identify the mechanism for this numerical difference. "
        "Malic acid was numerically similar (1.0 ± 0.16 g/L against 1.1 ± 0.23 g/L), so the summary data do not "
        "show a clear malolactic conversion pattern. However, no microbiological analyses, viable cell counts or "
        "method-sensitivity study were available. Residual microbial activity therefore cannot be ruled out. The "
        "result is reported as a descriptive compositional observation under the tested conditions, not as evidence "
        "that malolactic fermentation was absent."
    ),
    59: (
        "Volatile acidity, expressed as acetic acid, was numerically higher in the wine from the nominally warmer "
        "chamber (0.39 ± 0.03 g/L) than in the reference (0.30 ± 0.02 g/L) (Table 2). This magnitude is below "
        "commonly used spoilage limits for table wine, but the study did not include sensory testing. The "
        "descriptive between-condition difference is consistent with accelerated chemical ageing under warmer "
        "storage, including oxidative and hydrolytic pathways described for wine storage [20,21], but a "
        "microbiological contribution cannot be excluded because microbiological measurements were not collected."
    ),
    61: (
        "The descriptive summaries for pH, ethanol content and titratable acidity were similar between nominal "
        "chamber conditions. Total phenols, expressed as catechins, were also similar (1.02 ± 0.12 g/L in the "
        "reference against 1.10 ± 0.07 g/L in the nominally warmer chamber) (Table 2)."
    ),
    64: (
        "In accordance with the design described in Materials and Methods and summarised in Table 1, three stacks "
        "of 3-L BiB units were held for 20 days in each nominal chamber condition. The complete record is shown in "
        "Figure 3, with internal pressure above the temperature registered by the same sensors. In every usable "
        "trace, baseline-referred pressure showed a brief initial transient and then a broad rise-and-fall pattern. "
        "The traces also shared substantial common-mode behaviour, so the absolute pressure trajectory cannot be "
        "attributed only to package pressurisation and relaxation. Individual post-handling pressure maxima "
        "occurred between 85.0 and 138.5 h after the start of monitoring."
    ),
    66: (
        "After the handling window, uncorrected pressure summaries peaked between about days 3.5 and 5.8, at about "
        "56-94 mbar above each sensor's own baseline. Mean peak ΔP was 83.0 ± 17.2 mbar in bottom units and "
        "66.6 ± 10.1 mbar in top units under the nominal 50 °C chamber condition. Under the nominal 19 °C "
        "condition, the corresponding values were 84.8 ± 8.8 mbar and 59.6 ± 5.1 mbar (Figure 5A). In the "
        "exploratory sensor-level Type II two-way ANOVA, stack position was associated with uncorrected peak ΔP "
        "(F(1,8) = 9.14; p = 0.017). The adjusted bottom-minus-top difference was 20.4 mbar (95% CI 4.8-35.9 "
        "mbar). No statistically detectable nominal chamber-condition effect was observed (F(1,8) = 0.21; "
        "p = 0.66; adjusted nominal 50-minus-19 °C difference 3.1 mbar, 95% CI -12.5 to 18.6 mbar), and the "
        "interaction was not significant (p = 0.56). By day 20, neither position (F(1,8) = 0.35; p = 0.57; "
        "adjusted difference 5.2 mbar, 95% CI -15.3 to 25.7 mbar) nor nominal chamber condition (F(1,8) = 0.81; "
        "p = 0.39) was statistically detectable (Figure 5B). However, internal diagnostics showed strong common-mode "
        "structure across traces: leave-one-out correlations between each sensor and the mean of the remaining "
        "sensors averaged 0.94. When a global common-mode trace was subtracted and each sensor's maximum residual "
        "was analysed, the position term was not statistically detectable (F(1,8) = 3.21; p = 0.111; adjusted "
        "bottom-minus-top difference 14.4 mbar, 95% CI -4.1 to 32.9 mbar). Residuals at the original uncorrected "
        "peak times still showed a bottom-position offset (F(1,8) = 9.02; p = 0.017), and leave-one-out analysis "
        "of the uncorrected endpoint gave position p-values from 0.0007 to 0.051. The pressure result is therefore "
        "interpreted as an exploratory transient position-related offset superimposed on common-mode "
        "absolute-pressure variation."
    ),
    67: (
        "The lower panels of Figure 3 show that the realised thermal exposure differed substantially from the "
        "nominal set points. In the nominal 19 °C chamber, late in-bag temperatures averaged 23.50 ± 0.21 °C in "
        "top units and 22.84 ± 0.30 °C in bottom units. In the nominal 50 °C chamber, packages warmed gradually "
        "but remained far below the set point: 34.18 ± 0.36 °C at the top of the stacks and 26.08 ± 0.23 °C at "
        "the bottom during the late storage window. The warmest individual main-trial sensor reached 35.3 °C. By "
        "the end of storage, these temperatures were stable far below the nominal set point. This suggests "
        "sustained under-delivery and/or stratification of the effective thermal environment rather than only a "
        "short transient lag. Because no chamber-air logger was recorded, the chamber mechanism cannot be resolved. "
        "The pressure comparison between top and bottom positions must therefore be interpreted with the recognition "
        "that, in the nominally hot chamber, position also corresponded to different local package temperatures. "
        "The chamber comparison cannot be read as an actual 19 °C versus 50 °C wine-temperature contrast."
    ),
    68: (
        "Taken together, the physical and chemical measurements identify different response patterns, but they do "
        "not support a strict causal division between palletisation and temperature. The pressure measurements were "
        "position-resolved and showed transient position-related offsets in ΔP, superimposed on a common-mode "
        "absolute-pressure component. The chemical measurements were descriptive bulk summaries grouped by nominal "
        "chamber condition and were not resolved by stack position. The data support a practical distinction between "
        "the measurement types. Pressure monitoring identified position-related sensor offsets rather than a "
        "definitive mechanical-driver effect. Bulk chemical measurements indicated lower SO₂ and selected acidity "
        "differences under the nominally warmer chamber condition. The study does not show that palletisation and "
        "temperature act independently or exclusively on separate parts of the system."
    ),
    76: (
        "Figure 5. Change in internal pressure referred to each sensor's own baseline (ΔP), by stack position and "
        "nominal chamber condition. (A) post-handling peak value; individual sensor peaks occurred between 85.0 and "
        "138.5 h (about days 3.5-5.8). (B) residual value on day 20. Circles are individual sensors (n = 3 per "
        "position, except bottom position at nominal 19 °C, for which n = 2) and bars give mean ± 1 SD. At the "
        "uncorrected pressure peak, bottom-position BiBs had a larger ΔP than top-position BiBs in the exploratory "
        "sensor-level analysis. The two-way ANOVA gave F(1,8) = 9.14 and p = 0.017. An internal common-mode "
        "sensitivity analysis is reported in Section 3.2; the uncorrected p-value should be interpreted as "
        "exploratory rather than definitive. No statistically detectable nominal chamber-condition effect was "
        "observed (F(1,8) = 0.21; p = 0.66). By day 20 neither factor was statistically detectable."
    ),
    78: (
        "Table 2. Chemical parameters of wine from the main trial after 20 days in chambers nominally set at 19 °C "
        "and 50 °C. The labels refer to chamber set points, not actual in-bag wine temperatures. Values are "
        "descriptive summaries reported as mean ± standard deviation from the available study records. "
        "Replicate-level chemical records were not available to verify independent treatment-level replication; "
        "therefore, no inferential tests or significance letters are reported."
    ),
    85: (
        "The pressure response in the verification trial was small and transient (Figure 7A). Referred to each "
        "sensor's first valid pressure, internal pressure rose by 9.6-26.5 mbar early in the record. The traces "
        "then declined, showed a common negative excursion near day 6, and returned towards near-baseline values "
        "later in the record while the wine was held at about 45-48 °C. This pattern is consistent with flexible "
        "BiB package accommodation of thermal expansion, but the simultaneous movement of the traces also indicates "
        "a common-mode component. Because no barometric reference was logged, absolute ΔP traces should not be "
        "interpreted as package pressure alone. Because the verification trial used one stack and one configuration, "
        "it should be read as supporting evidence for that stack only. It does not show that temperature cannot "
        "affect BiB pressure under other configurations."
    ),
    90: (
        "This work was a preliminary, single-wine and single-format assessment, and its scope limits generalisation. "
        "The commercial wine's grape variety and vintage were unavailable. Only one BiB format was tested. Chemical "
        "analyses were grouped by nominal chamber condition and were not resolved by stack position, so the pressure "
        "and chemical datasets cannot be directly linked. Raw chemical replicate values were not retained in the "
        "available study records, and the triplicate determinations could not be verified as independent "
        "treatment-level replicates. Table 2 is therefore descriptive; SD values should not be interpreted as "
        "chamber-treatment replication. No microbiological measurements, viable cell counts, sensory analysis or "
        "post-transport storage follow-up were performed; therefore malolactic or other microbial contributions "
        "cannot be excluded and the consequences of lower SO₂ reserve remain potential implications rather than "
        "measured outcomes."
    ),
    91: (
        "The most important design limitation is the thermal mismatch in the main trial. The chamber nominally set "
        "to 50 °C produced actual late in-bag temperatures of 34.18 ± 0.36 °C at top sensors and 26.08 ± 0.23 °C "
        "at bottom sensors, with a persistent top-bottom gradient, so the 20-day comparison was not the intended "
        "19 °C versus 50 °C wine exposure. Chamber model, volume, fan power, airflow, chamber-air temperature and "
        "pallet occupancy were not recorded, preventing a definitive diagnosis of sustained under-delivery, "
        "stratification or heat-transfer limitation. There was also only one chamber per nominal set point, so "
        "chamber identity and nominal condition are confounded."
    ),
    92: (
        "The pressure analysis is limited by eleven usable traces, one missing bottom-position sensor, an unbalanced "
        "cell of n = 2, possible within-stack pairing of top and bottom observations, and a narrow stack-height "
        "range in which height was not independently replicated. The available raw workbook contains the "
        "pressure/temperature trajectories, but it does not retain stack IDs or sensor-to-stack pairings for a paired "
        "stack-level analysis. No barometric reference sensor was logged. Internal diagnostics showed strong "
        "common-mode correlations among pressure traces, and the uncorrected peak-position result was not robust to "
        "all endpoint definitions. The original peak endpoint was mostly stable in leave-one-out analysis, but one "
        "deletion gave p = 0.051; when each sensor's maximum common-mode residual was analysed, the position term "
        "was not statistically detectable. Baseline referencing reduces static sensor offsets, but it does not "
        "correct temperature-dependent pressure drift, ambient barometric variation, or possible perturbation from "
        "the oil-filled casing and displaced headspace. No bench calibration in the wine matrix was available. The "
        "first five hours were excluded from the primary storage analysis because they contained transport and "
        "stacking events. Sensitivity analysis showed that including this window did not change the pressure-peak "
        "endpoint. However, early transients may still be relevant to some logistics scenarios, and the available "
        "early workbook block contains duplicated rows. The verification trial approached higher in-bag temperatures "
        "than the main trial, but it used a single three-box stack for 15 days, one sensor stopped at day 4.35, and "
        "no sensor reached 50 °C. Finally, this study did not include a techno-economic analysis of sensor deployment, "
        "so monitoring recommendations should be understood as sentinel, validation or quality-control uses rather "
        "than universal instrumentation of every commercial package."
    ),
    93: (
        "Despite these limitations, the study illustrates why direct package-level measurement matters: chamber set "
        "points did not predict the realised wine temperature in the main trial, and sensor-level pressure summaries "
        "suggested position-related offsets under the tested stacking configuration. Future studies should include "
        "replicated chambers, chamber-air and barometric reference logging, bench calibration under relevant "
        "temperature and wine-matrix conditions, full release of original logger files, position-resolved chemical "
        "sampling, microbiological assays, realistic fluctuating temperature and vibration profiles, broader stack "
        "heights and formats, and cost-benefit evaluation of sentinel sensing strategies."
    ),
    95: (
        "This study monitored pressure and temperature inside 3-L Bag-in-Box wine units during static chamber storage "
        "intended to represent selected export-logistics stresses. Nominal chamber settings did not necessarily "
        "represent the temperature experienced by the wine. In the main trial, the chamber nominally set to 50 °C "
        "produced late in-bag temperatures of 34.18 ± 0.36 °C at the top of the stacks and 26.08 ± 0.23 °C at the "
        "bottom, while the reference chamber packages were about 23 °C. This mismatch changes the interpretation of "
        "the experiment: the main trial supports conclusions about the realised nominal chamber conditions and "
        "measured package temperatures, not about a 20-day exposure of wine to 50 °C."
    ),
    96: (
        "For pressure, bottom-position BiBs had a larger mean uncorrected peak ΔP than top-position BiBs under the "
        "tested stacking configuration, with an adjusted difference of 20.4 mbar in the exploratory sensor-level "
        "analysis (F(1,8) = 9.14; p = 0.017). Individual post-handling peaks occurred between about days 3.5 and "
        "5.8. However, the pressure traces shared a strong common-mode pattern and no barometric reference was "
        "logged. The position term for each sensor's maximum common-mode residual was not statistically detectable "
        "(F(1,8) = 3.21; p = 0.111), although residuals at the original uncorrected peak times remained higher for "
        "bottom-position sensors. The pressure result is therefore best interpreted as an exploratory transient "
        "position-related offset superimposed on common-mode absolute-pressure variation. Including the first five "
        "hours did not change the uncorrected peak-pressure endpoint. No statistically detectable nominal "
        "chamber-condition effect was observed in this small dataset, and by day 20 neither sensor-level factor was "
        "statistically detectable. The experiment does not provide a separate, well-powered test of stack height."
    ),
    97: (
        "For chemistry, descriptive summaries for wine from the nominally warmer chamber showed lower free and total "
        "SO₂, higher volatile acidity and a numerically higher lactic-acid value, while pH, ethanol, titratable "
        "acidity and total phenols were similar. Because the wine did not reach 50 °C in the main trial, chemical "
        "sampling was not position-resolved, and independent treatment-level replication could not be verified, "
        "these results should be interpreted as descriptive bulk compositional differences associated with the "
        "nominally warmer chamber condition. The lower SO₂ values may indicate reduced reserve available for later "
        "storage, but actual shelf life and sensory quality after transport were not measured. The verification "
        "trial, in which one three-box stack reached 45.0-48.8 °C across sensors without reaching 50 °C, showed no "
        "sustained pressure increase beyond transient changes of 9.6-26.5 mbar, with the same caveat about "
        "common-mode absolute-pressure components. Practical monitoring should therefore focus on selected sentinel "
        "packages or validation deployments that record actual in-package conditions and position-related pressure "
        "offsets, while economic feasibility and deployment density remain topics for future work."
    ),
    99: (
        "Author Contributions: Conceptualization, M.M., F.P., P.G.V. and A.Z.; methodology, B.C. and N.M.; software, "
        "B.C., K.M. and N.M.; validation, M.M., B.C. and N.M.; formal analysis, B.C., P.G.V. and N.M.; investigation, "
        "B.C. and N.M.; resources, F.P., P.G.V. and A.Z.; data curation, M.M., B.C., K.M. and N.M.; writing—original "
        "draft preparation, M.M. and N.M.; writing—review and editing, M.M. and N.M.; visualization, M.M., B.C. and "
        "N.M.; supervision, F.P. and A.Z.; project administration, F.P. and A.Z.; funding acquisition, A.Z."
    ),
    100: "All authors have read and agreed to the published version of the manuscript.",
    101: (
        "Funding: This research was funded by TRACEWINDU (Traceability at wine industry through integrated labelling "
        "of typicality, health protection effect and organoleptic attributes), a project co-funded by the Horizon "
        "2020 Framework Programme of the European Union under Grant Agreement no. 101007979, running from 1 June "
        "2021 to 31 May 2025."
    ),
    104: (
        "Data Availability Statement: The pressure and temperature workbooks, analysis scripts and derived summaries "
        "used for the revised analyses are available from the corresponding author and can be deposited as "
        "supplementary files if required by the journal. Raw chemical replicate records, microbiological data, "
        "sensory data, shelf-life follow-up data and chamber-control metadata were not retained in the available "
        "study records."
    ),
}


TABLE2_VALUES = [
    [
        "Sample",
        "Lactic Acid (g/L)",
        "Malic Acid (g/L)",
        "pH",
        "Titratable acidity (g/L of tartaric acid)",
        "Volatile acidity (g/L of acetic acid)",
        "Total phenols (g/L of catechins)",
        "Ethanol content (%v/v)",
        "Total Sulphur dioxide content\n(mg/L)",
        "Free Sulphur dioxide content\n(mg/L)",
    ],
    [
        "Wine from nominal 19 °C chamber",
        "0.16 ± 0.05",
        "1.1 ± 0.23",
        "3.4 ± 0.1",
        "6.6 ± 0.04",
        "0.30 ± 0.02",
        "1.02 ± 0.12",
        "11.6 ± 0.6",
        "50 ± 5.5",
        "31 ± 2.5",
    ],
    [
        "Wine from nominal 50 °C chamber",
        "0.29 ± 0.02",
        "1.0 ± 0.16",
        "3.4 ± 0.1",
        "6.6 ± 0.13",
        "0.39 ± 0.03",
        "1.10 ± 0.07",
        "11.6 ± 0.9",
        "25 ± 8.2",
        "16 ± 1.8",
    ],
]

TABLE2_NOTE = (
    "Values are descriptive summaries for wine sampled from each nominal chamber condition. Replicate-level "
    "chemical records were not available for independent treatment-level inference; therefore, no inferential "
    "tests or significance letters are reported. Where triplicate determinations represent analytical "
    "repeatability, SD should not be interpreted as chamber-treatment replication."
)


def set_paragraph_text(paragraph, text: str, highlighted: bool) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    if highlighted:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def highlight_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_cell_text(cell, text: str, highlighted: bool) -> None:
    cell.text = text
    if highlighted:
        highlight_cell(cell)


def delete_metadata_table(doc: Document) -> None:
    if doc.tables and "Academic Editor: Firstname Lastname" in doc.tables[0].cell(0, 0).text:
        table = doc.tables[0]._element
        table.getparent().remove(table)


def remove_last_column(table) -> None:
    xml_table = table._tbl
    grid = xml_table.tblGrid
    if grid is not None and len(grid.gridCol_lst) > 10:
        grid.remove(grid.gridCol_lst[-1])
    for row in xml_table.tr_lst:
        cells = list(row.tc_lst)
        if len(cells) > 10:
            row.remove(cells[-1])
        elif len(cells) == 1:
            tc_pr = cells[0].tcPr
            if tc_pr is not None and tc_pr.gridSpan is not None:
                try:
                    span = int(tc_pr.gridSpan.val)
                except (TypeError, ValueError):
                    span = 11
                if span > 10:
                    tc_pr.gridSpan.val = 10


def find_table2(doc: Document):
    for table in doc.tables:
        if table.cell(0, 0).text.strip() == "Sample":
            return table
    raise RuntimeError("Table 2 not found")


def update_table2(doc: Document, highlighted: bool) -> None:
    table = find_table2(doc)
    if len(table.columns) > 10:
        remove_last_column(table)
        table = find_table2(doc)
    for row_idx, values in enumerate(TABLE2_VALUES):
        for col_idx, value in enumerate(values):
            set_cell_text(
                table.cell(row_idx, col_idx),
                value,
                highlighted and row_idx > 0 and col_idx > 1,
            )
    set_cell_text(table.cell(3, 0), TABLE2_NOTE, highlighted)


def update_doc(path: Path, highlighted: bool) -> None:
    doc = Document(path)
    for idx, text in PARAGRAPH_UPDATES.items():
        if idx >= len(doc.paragraphs):
            raise RuntimeError(f"{path}: missing paragraph {idx}")
        set_paragraph_text(doc.paragraphs[idx], text, highlighted)
    delete_metadata_table(doc)
    update_table2(doc, highlighted)
    doc.save(path)


def main() -> None:
    update_doc(ROOT_CLEAN, highlighted=False)
    update_doc(ROOT_HIGHLIGHTED, highlighted=True)
    copy2(ROOT_HIGHLIGHTED, PKG_HIGHLIGHTED)
    copy2(ROOT_CLEAN, PKG_CLEAN)
    print("Updated final clean/highlighted manuscripts and submission copies.")


if __name__ == "__main__":
    main()

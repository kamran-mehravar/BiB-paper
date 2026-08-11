#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Synchronise response-letter wording after the go/no-go audit."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


TITLE_OLD = "Palletisation and In-Bag Thermal Exposure in Bag-in-Box Wine Packaging under Simulated Export Conditions"
TITLE_NEW = "In-Bag Pressure and Temperature Monitoring of Palletised 3-L Bag-in-Box Wine under Static Chamber Conditions"

MD_FILES = [
    Path("Response_to_Reviewers_R2_R3.md"),
    Path("SUBMISSION_PACKAGE/Responses/03_Response_to_Reviewer_2.md"),
    Path("SUBMISSION_PACKAGE/Responses/04_Response_to_Reviewer_3.md"),
]

REPLACEMENTS = {
    TITLE_OLD: TITLE_NEW,
    (
        "This included re-analysis of the main-trial pressure and temperature time-series, a sensitivity analysis "
        "of the initial handling period, a more detailed examination of the verification trial, and revision of "
        "Figures 3-7 to ensure that nominal chamber set points are clearly distinguished from measured in-bag "
        "temperatures."
    ): (
        "This included re-analysis of the main-trial pressure and temperature time-series, an internal common-mode "
        "pressure diagnostic, a sensitivity analysis of the initial handling period, a descriptive-only presentation "
        "of Table 2 chemistry, a more detailed examination of the verification trial, and revision of Figures 3-7 to "
        "ensure that nominal chamber set points are clearly distinguished from measured in-bag temperatures."
    ),
    (
        "The Abstract has been revised to reflect the re-analysis of the original experimental records. It now "
        "reports the late-temperature summaries, the exploratory pressure effect estimate, the absence of a "
        "statistically detectable nominal chamber-condition effect, the summary-only scope of the chemistry data, "
        "and the limited role of the verification trial."
    ): (
        "The Abstract has been revised to reflect the re-analysis of the original experimental records. It now "
        "reports the late-temperature summaries, the uncorrected exploratory pressure result, the common-mode "
        "pressure sensitivity, the descriptive-only scope of the chemistry data, and the limited role of the "
        "verification trial."
    ),
    (
        "We agree. The re-analysis improves reproducibility but does not change the replication status. We "
        "reanalysed the original main-trial sensor records and independently reconstructed the pressure summaries "
        "used in the statistical analysis. The reconstructed values match the reported summaries to 0.0000 mbar. "
        "However, the pressure cells remain n = 3, 3, 3 and 2, chamber condition remains one chamber per nominal "
        "set point, and the records do not retain stack IDs or sensor-to-stack pairings for a blocked or paired "
        "stack-level analysis. We therefore retain the ANOVA only as exploratory sensor-level inference and report "
        "effect estimates and confidence intervals alongside p-values."
    ): (
        "We agree. The re-analysis improves reproducibility but does not change the replication status. We "
        "reanalysed the original main-trial sensor records and independently reconstructed the pressure summaries "
        "used in the statistical analysis. The reconstructed values match the reported summaries to 0.0000 mbar. "
        "However, the pressure cells remain n = 3, 3, 3 and 2, chamber condition remains one chamber per nominal "
        "set point, and the records do not retain stack IDs or sensor-to-stack pairings for a blocked or paired "
        "stack-level analysis. We also added an internal common-mode diagnostic because no barometric reference "
        "sensor was logged. The uncorrected peak-position result remains F(1,8) = 9.14 and p = 0.017, but the "
        "position term for each sensor's maximum common-mode residual was not statistically detectable "
        "(F(1,8) = 3.21; p = 0.111). We therefore present the pressure result as an exploratory transient "
        "position-related offset rather than a definitive package-mechanics effect."
    ),
    (
        "Sections 2.2, 2.4, 3.2, 3.4 and Conclusions now state the experimental unit, the missing sensor, the "
        "unbalanced design, the unreplicated chamber factor and the exploratory interpretation. The analysis report "
        "and `anova_pressure.py` were updated to document the raw-workbook reconstruction."
    ): (
        "Sections 2.2, 2.4, 3.2, 3.4 and Conclusions now state the experimental unit, the missing sensor, the "
        "unbalanced design, the unreplicated chamber factor, the absence of barometric reference logging and the "
        "exploratory interpretation. The analysis report and reproducibility scripts document the raw-workbook "
        "reconstruction and additional common-mode pressure diagnostics."
    ),
    (
        "We agree. The manuscript no longer presents this binary causal attribution. The original study records do "
        "not contain position-resolved chemistry. The final text states that pressure was measured by package "
        "position and showed a transient bottom-position peak, whereas chemistry was measured as bulk wine grouped "
        "by nominal chamber condition. These two datasets support different descriptive response patterns, but they "
        "do not demonstrate that palletisation and temperature act independently or exclusively on separate parts "
        "of the system."
    ): (
        "We agree. The manuscript no longer presents this binary causal attribution. The original study records do "
        "not contain position-resolved chemistry. The final text states that pressure was measured by package "
        "position and showed transient position-related offsets superimposed on a common-mode absolute-pressure "
        "component, whereas chemistry is presented as descriptive bulk summaries grouped by nominal chamber "
        "condition. These two datasets support different descriptive response patterns, but they do not demonstrate "
        "that palletisation and temperature act independently or exclusively on separate parts of the system."
    ),
    (
        "The Abstract, Discussion, Limitations and Conclusions were revised; unsupported \"governs\" language was "
        "removed."
    ): (
        "The Abstract, Discussion, Limitations and Conclusions were revised; unsupported \"governs\" language was "
        "removed and the pressure interpretation was further limited by the common-mode diagnostic."
    ),
    (
        "We agree. We re-examined the available original experimental records and confirmed that raw replicate-level "
        "chemical measurements were not retained in the available study records. The originally submitted Table 2 "
        "reports `0.29 +/- 0.02 g/L`, whereas no primary experimental record supports the alternative "
        "`0.29 +/- 0.20 g/L`. The manuscript therefore reports `0.29 +/- 0.02 g/L` as the traceable descriptive "
        "Table 2 value, without a significance claim. We state that malic acid did not show the clear decrease "
        "expected for a simple malolactic-conversion interpretation, but we no longer state that malolactic "
        "fermentation was ruled out. Residual microbial activity cannot be excluded."
    ): (
        "We agree. We re-examined the available original experimental records and confirmed that raw replicate-level "
        "chemical measurements were not retained in the available study records. The originally submitted Table 2 "
        "reports `0.29 +/- 0.02 g/L`, whereas no primary experimental record supports the alternative "
        "`0.29 +/- 0.20 g/L`. The manuscript therefore reports `0.29 +/- 0.02 g/L` as the traceable descriptive "
        "Table 2 value, without a significance claim. Because independent treatment-level replication cannot be "
        "verified for the chemistry table, Tukey letters and chemistry p-values have been removed. Malic acid was "
        "numerically similar between conditions, but we no longer state that malolactic fermentation was ruled out. "
        "Residual microbial activity cannot be excluded."
    ),
    (
        "Sections 2.4 and 3.1, Table 2 and Limitations were revised to report lactic acid descriptively using the "
        "traceable `0.29 +/- 0.02 g/L` value and to state the absence of raw chemical replicates, microbiology, "
        "viable cell counts and method-sensitivity data."
    ): (
        "Sections 2.1, 2.4 and 3.1, Table 2 and Limitations were revised to report chemistry descriptively using "
        "the traceable Table 2 values and to state the absence of raw chemical replicates, microbiology, viable cell "
        "counts and method-sensitivity data."
    ),
    (
        "Corrected. Table 2 now labels the rows as wine from the nominal 19 degC chamber and wine from the nominal "
        "50 degC chamber. The caption states that these labels are chamber set points, not actual in-bag wine "
        "temperatures."
    ): (
        "Corrected. Table 2 now labels the rows as wine from the nominal 19 degC chamber and wine from the nominal "
        "50 degC chamber. The caption states that these labels are chamber set points, not actual in-bag wine "
        "temperatures. Because independent treatment-level replication could not be verified for the chemistry "
        "records, the table now reports descriptive summaries only and contains no significance letters."
    ),
    "Table 2 row labels, caption and note were revised and highlighted.": (
        "Table 2 row labels, caption, values and note were revised and highlighted; significance letters were removed."
    ),
    (
        "We appreciate this point. The final manuscript clarifies that the pressure analysis evaluates top versus "
        "bottom package position under the tested stack configuration. Stack height was not analysed as an "
        "independent factor because the design contained one four-box stack and two three-box stacks per nominal "
        "chamber condition. Re-examination of the original main-trial records confirmed the available sensor traces "
        "and position/chamber mapping; however, stack identifiers or sensor-to-stack pairings were not retained. "
        "Chamber condition also remains unreplicated at the chamber level. The pressure ANOVA is therefore presented "
        "as exploratory sensor-level analysis rather than a definitive test of stack height, chamber temperature or "
        "paired stack effects."
    ): (
        "We appreciate this point. The final manuscript clarifies that the pressure analysis evaluates top versus "
        "bottom package position under the tested stack configuration. Stack height was not analysed as an "
        "independent factor because the design contained one four-box stack and two three-box stacks per nominal "
        "chamber condition. Re-examination of the original main-trial records confirmed the available sensor traces "
        "and position/chamber mapping; however, stack identifiers or sensor-to-stack pairings were not retained. "
        "Chamber condition also remains unreplicated at the chamber level, and no barometric reference was logged. "
        "The pressure ANOVA is therefore presented as exploratory sensor-level analysis rather than a definitive "
        "test of stack height, chamber temperature, package mechanics or paired stack effects."
    ),
    (
        "We reanalysed the original temperature records to address this point quantitatively. Late in the main "
        "trial, nominal 50 degC packages averaged 34.18 +/- 0.36 degC at top sensors and 26.08 +/- 0.23 degC at "
        "bottom sensors; nominal 19 degC packages averaged 23.50 +/- 0.21 degC at top sensors and 22.84 +/- 0.30 "
        "degC at bottom sensors. The manuscript now separates observed facts from interpretation. The observations "
        "are consistent with thermal inertia, chamber loading, heat-transfer resistance through the pallet/stack "
        "and airflow constraints, but the chamber model, airflow and occupancy data were not recorded, so these "
        "remain hypotheses rather than established causes."
    ): (
        "We reanalysed the original temperature records to address this point quantitatively. Late in the main "
        "trial, nominal 50 degC packages averaged 34.18 +/- 0.36 degC at top sensors and 26.08 +/- 0.23 degC at "
        "bottom sensors; nominal 19 degC packages averaged 23.50 +/- 0.21 degC at top sensors and 22.84 +/- 0.30 "
        "degC at bottom sensors. The manuscript now separates observed facts from interpretation. By the end of "
        "storage, the main-trial temperatures were stable far below the nominal 50 degC set point, suggesting "
        "sustained under-delivery and/or stratification of the effective thermal environment rather than only a "
        "short transient lag. Because no chamber-air logger, chamber model, airflow or occupancy data were recorded, "
        "the chamber mechanism cannot be resolved."
    ),
    (
        "Sections 3.2, 3.3, 3.4 and the captions to Figures 3 and 6 were revised with the quantified temperatures "
        "and cautious interpretation."
    ): (
        "Sections 3.2, 3.3, 3.4 and the captions to Figures 3 and 6 were revised with the quantified temperatures "
        "and the sustained-under-delivery/stratification limitation."
    ),
}


DOCX_TARGETS = {
    Path("Response_to_Reviewers_R2_R3.md"): Path("Response_to_Reviewers.docx"),
    Path("SUBMISSION_PACKAGE/Responses/03_Response_to_Reviewer_2.md"): Path(
        "SUBMISSION_PACKAGE/Responses/03_Response_to_Reviewer_2.docx"
    ),
    Path("SUBMISSION_PACKAGE/Responses/04_Response_to_Reviewer_3.md"): Path(
        "SUBMISSION_PACKAGE/Responses/04_Response_to_Reviewer_3.docx"
    ),
}


def update_markdown(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    path.write_text(text, encoding="utf-8")


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(72)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            paragraph = doc.add_paragraph()
            if line.startswith("**Reviewer Comment**") or line.startswith("**Response**") or line.startswith(
                "**Changes in Manuscript**"
            ):
                run = paragraph.add_run(line.replace("**", ""))
                run.bold = True
            else:
                paragraph.add_run(line.replace("**", ""))
    doc.save(docx_path)


def main() -> None:
    for path in MD_FILES:
        update_markdown(path)
    for md_path, docx_path in DOCX_TARGETS.items():
        markdown_to_docx(md_path, docx_path)
    print("Updated reviewer-response Markdown and DOCX files.")


if __name__ == "__main__":
    main()
